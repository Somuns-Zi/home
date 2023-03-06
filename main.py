import os
import re
import shutil
import time

from docx import Document  # 需下载python-docx-2023
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
from openpyxl import load_workbook  # 需下载
from win32com import client     #下载pywin32即可

# 大循环前的初始化=========================================

def homework_correcting():
    marking_date = input("请输入阅卷日期：\n示例：2022.6.13\n")
    marking_teacher = input("请输入阅卷老师：\n示例：张永锐\n")

    while True:
        os.system("cls")
        krow = 1  # 定义写入c成绩表格的初始行号
        kcolumn = 3  # 定义写入成绩表格的初始列号
        class_name = input("请输入班级名：\n示例：20190719\n")
        if (os.path.exists(f'd:\减速器拆装与测量实验报告批改\减速器拆装与测量实验报告批改工作区\{class_name}')):
            wordfilepath = f'd:\减速器拆装与测量实验报告批改\减速器拆装与测量实验报告批改工作区\{class_name}'  # word文件夹路径
        else:
            while (True):
                class_name = input("请输入正确的班级名称!!!!\n:")
                if (os.path.exists(f'd:\减速器拆装与测量实验报告批改\减速器拆装与测量实验报告批改工作区\{class_name}')):
                    wordfilepath = f'd:\减速器拆装与测量实验报告批改\减速器拆装与测量实验报告批改工作区\{class_name}'  # word文件夹路径
                    break
        num = int(class_name)
        num_one = int(num / 10000)
        num_two = int(num_one * 100)
        num_three = int(num / 100)
        faculty = num_three - num_two
        if (
        os.path.exists(f"D:\减速器拆装与测量实验报告批改\批完减速器拆装与测量实验报告作业\减速器批完报告\{faculty}系")):
            doc_to_docx_one_path = f"D:\减速器拆装与测量实验报告批改\批完减速器拆装与测量实验报告作业\减速器批完报告\{faculty}系"
        else:
            doc_to_docx_one_path = os.mkdir(
                f"D:\减速器拆装与测量实验报告批改\批完减速器拆装与测量实验报告作业\减速器批完报告\{faculty}系")
        riqi = input(
            f"请输入{class_name}班的上课日期：\n示例：2022.5.15\n")  # 实验日期缺省值！！！！！！！！！！！！！！！！！！！！！！！！！！！学生实验日期，未填写情况下帮忙填写！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！
        yichangfile = 0  # 未处理的非doc/docx类型的实验报告数量
        if (os.path.exists(
                f"D:\减速器拆装与测量实验报告批改\批完减速器拆装与测量实验报告作业\减速器批完报告\{faculty}系\{class_name}")):
            doc_to_docx_path = f"D:\减速器拆装与测量实验报告批改\批完减速器拆装与测量实验报告作业\减速器批完报告\{faculty}系\{class_name}"
        else:
            doc_to_docx_two = os.mkdir(
                f"D:\减速器拆装与测量实验报告批改\批完减速器拆装与测量实验报告作业\减速器批完报告\{faculty}系\{class_name}")
            doc_to_docx_path = f"D:\减速器拆装与测量实验报告批改\批完减速器拆装与测量实验报告作业\减速器批完报告\{faculty}系\{class_name}"

        ##================================================
        def doc_to_docx(path_original, path_final):  # doc转docx函数
            if os.path.splitext(path_original)[1] == ".doc":
                word = client.Dispatch('Word.Application')
                print(path_original)
                doc = word.Documents.Open(path_original)  # 目标路径下的文件
                print(path_final)
                doc.SaveAs(path_final, 16)  # 转化后路径下的文件
                doc.Close()
                # word.Quit()
            elif os.path.splitext(path_original)[1] == ".docx":
                shutil.copy(path_original, path_final)

        ##================================================
        def get_pictures(word_path, result_path):  # 获取word内图片的函数
            doc = Document(word_path)
            dict_rel = doc.part._rels
            pic_num = 0
            for rel in dict_rel:
                try:
                    rel = dict_rel[rel]
                    print(rel)
                    if "image" in rel.target_ref:
                        pic_num += 1
                        if not os.path.exists(result_path):
                            os.makedirs(result_path)
                        img_name = re.findall("/(.*)", rel.target_ref)[0]
                        word_name = os.path.splitext(word_path)[0]
                        # print(os.sep)
                        if os.sep in word_name:
                            new_name = word_name.split('\\')[-1]
                        else:
                            new_name = word_name.split('/')[-1]
                        img_name = f'{img_name}'
                        with open(f'{result_path}/{img_name}', "wb") as f:
                            f.write(rel.target_part.blob)
                except:
                    continue
            return (pic_num)

        ##================================================
        def del_file(path):  # 删文件函数
            ls = os.listdir(path)
            for i in ls:
                c_path = os.path.join(path, i)
                if os.path.isdir(c_path):
                    del_file(c_path)
                else:
                    os.remove(c_path)

        ##================================================
        def createfile(b):  # 创建文件夹函数
            if (not (os.path.exists(b))):
                os.mkdir(b)

        ##================定义表格题的三类减速器答案范围================================
        def check_1(x):
            y = 20
            try:
                sta = [[19.2, 20.2], [11.4, 11.9], [9.2, 9.9], [6.9, 7.7], [5.4, 5.9]
                    , [7.4, 8.4], [15.2, 17.2], [14.1, 16.9], [22.1, 22.8], [113.5, 115.8]
                    , [138.1, 139.9], [153.1, 154.9], [7.3, 7.9], [7.4, 8.1], [19.1, 22.1]
                    , [4.3, 8.6]]
                for i in range(1, 17, 1):
                    if len(x[i]) == 0:
                        y = y - 2
                    else:
                        if float(x[i]) < sta[i - 1][0] or float(x[i]) > sta[i - 1][1]:
                            y = y - 2
            except:
                y = 14

            return y

        def check_2(x):
            y = 20
            try:
                sta = [[19.2, 20.2], [11.4, 11.9], [9.2, 9.9], [6.9, 7.7], [5.4, 5.9]
                    , [7.4, 8.9], [15.2, 17.2], [14.1, 16.9], [22.1, 22.8], [200, 203]
                    , [0, 0], [201, 204.5], [7.3, 7.9], [7.4, 8.3], [18.1, 22.1], [4.3, 8.6]]
                for i in range(1, 17, 1):
                    if i == 11:
                        if len(x[i]) > 0:
                            y = y - 2
                    elif len(x[i]) > 0:
                        if float(x[i]) < sta[i - 1][0] or float(x[i]) > sta[i - 1][1]:
                            y = y - 2
                    else:
                        y = y - 2
            except:
                y = 14

            return y

        def check_3(x):
            y = 20
            try:
                sta = [[19.2, 20.2], [11.4, 11.9], [9.2, 9.9], [6.9, 7.7], [5.4, 5.9]
                    , [7.4, 8.9], [15.2, 17.2], [14.1, 16.9], [22.1, 22.8], [140.1, 142.5]
                    , [0, 0], [184.1, 185.8], [7.4, 8.3], [7.4, 8.3], [22.1, 24.5], [4.3, 8.6]]
                for i in range(1, 17, 1):
                    if i == 11:
                        if len(x[i]) > 0:
                            y = y - 2
                    elif len(x[i]) > 0:
                        if float(x[i]) < sta[i - 1][0] or float(x[i]) > sta[i - 1][1]:
                            y = y - 2
                    else:
                        y = y - 2
            except:
                y = 14

            return y

        ##================================================

        pigaifilepath = doc_to_docx_path  # 批改保存路径
        a = os.path.dirname(wordfilepath)

        b = os.path.join(a, '图片')
        createfile(b)
        result_path = b  # 图片保存路径

        excelfilepath = r'd:\减速器拆装与测量实验报告批改\减速器拆装与测量实验报告批改工作区\小分.xlsx'  # ！！！！！！！！！！！！！！！！！！！！！！！！！样本成绩表路径！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！！
        new_excelpath = os.path.dirname(wordfilepath)  # 表格保存母路径
        newexcel = os.path.join(new_excelpath, '本班成绩表.xlsx')  # 表格保存全路径
        shutil.copy(excelfilepath, newexcel)

        ##======================更改doc为docx==========================

        spam = os.listdir(wordfilepath)  # 获取文件夹下的word文档列表
        ms = 0  # 记录更改失败次数
        for word in spam:
            try:
                path_original = os.path.join(wordfilepath, word)  # 连接为doc的完整word路径
                if word.endswith(".xlsx"):
                    shutil.move(path_original, wordfilepath + f"\\{class_name}.xlsx")
                if word.endswith(".doc"):
                    path_final = os.path.join(doc_to_docx_path, word) + 'x'  # 连接为docx的完整word路径
                elif word.endswith('.docx'):
                    path_final = os.path.join(doc_to_docx_path, word)
                else:
                    yichangfile += 1
                    continue
                doc_to_docx(path_original, path_final)
            except:
                ms += 1  # 更改失败次数

        print('本班级后缀异常的实验报告有 ' + str(yichangfile) + ' 个' + '，转格式失败的实验报告有 ' + str(
            ms) + ' 个' + '，请核查' + '\n')
        input("按任意键继续")
        os.system("cls")

        ##======================改名===================================
        spam = os.listdir(pigaifilepath)  # 获取文件夹下的word文档列表
        error = 0
        for word in spam:
            try:
                if word.endswith(".docx"):
                    filename = os.path.join(pigaifilepath, word)  # 连接为完整的word路径
                else:

                    continue
                document = Document(filename)  # 打开指定目录下的word文档

                paragraphs = document.paragraphs  # 获得文档的所有段落
                length = len(paragraphs)  # 获得文档总段落数
                tb = document.tables
                cell = tb[0].rows[4].cells[1]
                name = cell.text
                cell = tb[0].rows[3].cells[4]
                ID = cell.text
                shutil.move(filename, pigaifilepath + "\\" + ID + name + "减速器拆装与测量实验报告.docx")
            except:
                error += 1
        print(f"改名失败的文档有{error}个。")
        input("按下任意键继续")
        os.system("cls")
        ##======================开始大循环================================

        spam = os.listdir(pigaifilepath)  # 获取文件夹下的word文档列表
        k = 0  # 用来表示是第几个循环
        yichang = 0  # 异常报告内容数量
        for word in spam:
            try:
                krow += 1  # 成绩表打印索引
                k += 1  # 循环索引
                if word.endswith(".docx"):
                    filename = os.path.join(pigaifilepath, word)  # 连接为完整的word路径
                else:

                    continue
                ##===============大循环内的初始化======================

                i = 0  # 所打开word文档的段落号索引，大循环内

                ##=====================================================

                wb = load_workbook(newexcel)  # 打开指定目录下的表格
                ws = wb.active  # 激活表格

                document = Document(filename)  # 打开指定目录下的word文档

                paragraphs = document.paragraphs  # 获得文档的所有段落
                length = len(paragraphs)  # 获得文档总段落数
                ##======================获取图片==============================
                print(f'正在获取第{k}个报告的图片。。。')
                del_file(result_path)  # 先删除原图片
                pic_num = get_pictures(filename, result_path)  # 添加现图片

                ##============================================================
                # 检查、帮忙填入实验时间、指导教师、教师签字、日期
                tb = document.tables
                cell = tb[0].rows[5].cells[4]
                value = cell.text
                run = cell.paragraphs[0]
                run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if len(value.strip()) > 0:
                    riqi = value
                else:
                    run1 = run.add_run(riqi)
                    run1.font.size = Pt(12)
                    run1.font.color.rgb = RGBColor(0, 0, 0)
                ##---------------------------------
                cell = tb[0].rows[4].cells[4]
                value = cell.text
                run = cell.paragraphs[0]
                run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.clear()
                run1 = run.add_run(marking_teacher)
                run1.font.size = Pt(12)
                run1.font.color.rgb = RGBColor(0, 0, 0)
                ##---------------------------------
                cell = tb[0].rows[8].cells[1]
                value = cell.text
                run = cell.paragraphs[0]
                run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.clear()
                run1 = run.add_run(marking_teacher)
                run1.font.size = Pt(12)
                run1.font.color.rgb = RGBColor(255, 0, 0)
                ##---------------------------------
                cell = tb[0].rows[8].cells[4]
                value = cell.text
                run = cell.paragraphs[0]
                run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.clear()
                run1 = run.add_run(marking_date)  # ！！！！！！！！！！！！！！！！！！！！！！！！！！！！这里修改批改日期！！！！！！！！！！！！！！！！！！！！！！！！！！！
                run1.font.size = Pt(12)
                run1.font.color.rgb = RGBColor(255, 0, 0)

                cell = tb[0].rows[4].cells[1]
                name = cell.text
                cell = tb[0].rows[3].cells[4]
                ID = cell.text

                # ==============================获取学生作答内容==============================================
                n = 0  # 检查是否有题目缺漏
                i1_lock = 0
                i2_lock = 0
                i3_lock = 0
                i4_lock = 0
                i5_lock = 0
                i6_lock = 0
                i7_lock = 0
                i8_lock = 0
                i9_lock = 0
                i10_lock = 0
                for par in document.paragraphs:

                    if "说明减速器的组成及主要用途" in par.text and i1_lock == 0:
                        i1 = i
                        paragraphs[i1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        n += 1
                        i1_lock = 1

                    if "实验目的" in par.text and i2_lock == 0:
                        i2 = i
                        paragraphs[i2].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        n += 1
                        i2_lock = 1

                    if "实验设备及主要工具" in par.text and i3_lock == 0:
                        i3 = i
                        paragraphs[i3].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        n += 1
                        i3_lock = 1

                    if "实验拆装步骤" in par.text and i4_lock == 0:
                        i4 = i
                        paragraphs[i4].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        n += 1
                        i4_lock = 1

                    if "注意事项" in par.text and i5_lock == 0:
                        i5 = i
                        paragraphs[i5].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        n += 1
                        i5_lock = 1

                    if "画出你所拆装的减速器" in par.text and i6_lock == 0:
                        i6 = i
                        paragraphs[i6].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        n += 1
                        i6_lock = 1

                    if "将测量数据填入" in par.text and i7_lock == 0:
                        i7 = i
                        paragraphs[i7].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        n += 1
                        i7_lock = 1

                    if "说明减速器各零件" in par.text and i8_lock == 0:
                        i8 = i
                        paragraphs[i8].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        n += 1
                        i8_lock = 1

                    if "你所拆装的减速器采用的" in par.text and i9_lock == 0:
                        i9 = i
                        paragraphs[i9].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        n += 1
                        i9_lock = 1

                    if "箱体与箱盖接触面是如何解决密封的" in par.text and i10_lock == 0:
                        i10 = i
                        paragraphs[i10].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        n += 1
                        i10_lock = 1

                    i = i + 1
                ilist = [i1, i2, i3, i4, i5, i6, i7, i8, i9, i10]
                lengthi = len(ilist)
                i11 = length
                i22 = length
                i33 = length
                i44 = length
                i55 = length
                i66 = length
                i77 = length
                i88 = length
                i99 = length
                i1010 = length
                for ki in range(lengthi):
                    if ilist[ki] > i1 and ilist[ki] < i11:
                        i11 = ilist[ki]

                    if ilist[ki] > i2 and ilist[ki] < i22:
                        i22 = ilist[ki]

                    if ilist[ki] > i3 and ilist[ki] < i33:
                        i33 = ilist[ki]

                    if ilist[ki] > i4 and ilist[ki] < i44:
                        i44 = ilist[ki]

                    if ilist[ki] > i5 and ilist[ki] < i55:
                        i55 = ilist[ki]

                    if ilist[ki] > i6 and ilist[ki] < i66:
                        i66 = ilist[ki]

                    if ilist[ki] > i7 and ilist[ki] < i77:
                        i77 = ilist[ki]

                    if ilist[ki] > i8 and ilist[ki] < i88:
                        i88 = ilist[ki]

                    if ilist[ki] > i9 and ilist[ki] < i99:
                        i99 = ilist[ki]

                    if ilist[ki] > i10 and ilist[ki] < i1010:
                        i1010 = ilist[ki]
                if (n == 10):
                    answer1 = ''
                    for j in range(i1 + 1, i11, 1):
                        answer1 = answer1 + '\n' + str(paragraphs[j].text)

                    answer2 = ''
                    for j in range(i2 + 1, i22, 1):
                        answer2 = answer2 + '\n' + str(paragraphs[j].text)

                    answer3 = ''
                    for j in range(i3 + 1, i33, 1):
                        answer3 = answer3 + '\n' + str(paragraphs[j].text)

                    answer4 = ''
                    for j in range(i4 + 1, i44, 1):
                        answer4 = answer4 + '\n' + str(paragraphs[j].text)

                    answer5 = ''
                    for j in range(i5 + 1, i55, 1):
                        answer5 = answer5 + '\n' + str(paragraphs[j].text)

                    answer6 = ''
                    for j in range(i6 + 1, i66, 1):
                        answer6 = answer6 + '\n' + str(paragraphs[j].text)

                    answer9 = ''
                    for j in range(i9 + 1, i99, 1):
                        answer9 = answer9 + '\n' + str(paragraphs[j].text)

                    answer10 = ''
                    for j in range(i10 + 1, i1010, 1):
                        answer10 = answer10 + '\n' + str(paragraphs[j].text)

                    ##==============================================================================================
                    cell = tb[0].rows[4].cells[1]
                    name = cell.text
                    cell = tb[0].rows[3].cells[4]
                    ID = cell.text
                    os.system("cls")

                    # =================输入第一题（需要阅览的答案）得分===================
                    print(f'学生{k}：' + name + ID + '，共有图片' + str(pic_num) + '张')
                    if len(answer1.strip()) > 1:
                        print("第 1 题回答是：" + answer1.strip())
                    else:
                        print("第 1 题回答可能是图片类型：" + answer1.strip())
                    print('\n''标准答案为：*************')
                    score1 = input("请输入第1题得分（3分）:")

                    while ((not (score1.isdigit())) or (int(score1) > 3) or (int(score1) < 0)):
                        score1 = input('请输入分值范围内的数字！第1题得分（3分）:')
                    print('第1题得分为：' + score1)
                    os.system("cls")

                    print(f'学生{k}：' + name + ID + '，共有图片' + str(pic_num) + '张')
                    # ===============判断第二至五题（不需要阅览回答内容的题目）得分==============
                    if len(answer2.strip()) > 10:
                        score2 = str(3)
                    else:
                        print("第 2 题回答是：可能是图片类型：" + answer2.strip())
                        print('\n''标准答案为：*************')
                        score2 = input('请输入第2题得分（3分）:')
                        while ((not (score2.isdigit())) or (int(score2) > 3) or (int(score2) < 0)):
                            score2 = input('请输入分值范围内的数字！第2题得分（3分）:')
                        print('第2题得分为：' + score2)

                        print(f'学生{k}：' + name + ID + '，共有图片' + str(pic_num) + '张')
                    os.system("cls")
                    if len(answer3.strip()) > 10:
                        score3 = str(3)
                    else:
                        print("第 3 题回答是：可能是图片类型：" + answer3.strip())
                        print('\n''标准答案为：*************')
                        score3 = input('请输入第3题得分（3分）:')
                        while ((not (score3.isdigit())) or (int(score3) > 3) or (int(score3) < 0)):
                            score3 = input('请输入分值范围内的数字！第3题得分（3分）:')
                        print('第3题得分为：' + score3)

                        print(f'学生{k}：' + name + ID + '，共有图片' + str(pic_num) + '张')
                    os.system("cls")
                    if len(answer4.strip()) > 10:
                        score4 = str(3)
                    else:
                        print("第 4 题回答是：可能是图片类型：" + answer4.strip())
                        print('\n''标准答案为：*************')
                        score4 = input('请输入第4题得分（3分）:')
                        while ((not (score4.isdigit())) or (int(score4) > 3) or (int(score4) < 0)):
                            score4 = input('请输入分值范围内的数字！第4题得分（3分）:')
                        print('第4题得分为：' + score4)

                        print(f'学生{k}：' + name + ID + '，共有图片' + str(pic_num) + '张')
                    os.system("cls")
                    if len(answer5.strip()) > 10:
                        score5 = str(3)
                    else:
                        print("第 5 题回答是：可能是图片类型：" + answer5.strip())
                        print('\n''标准答案为：*************')
                        score5 = input('请输入第5题得分（3分）:')
                        while ((not (score5.isdigit())) or (int(score5) > 3) or (int(score5) < 0)):
                            score5 = input('请输入分值范围内的数字！第5题得分（3分）:')
                        print('第5题得分为：' + score4)
                        os.system("cls")

                    print(f'学生{k}：' + name + ID + '，共有图片' + str(pic_num) + '张')
                    ##===========================给出第六题得分==============================================
                    print("第 6 题回答是(应该是图片类型:)" + answer6.strip())
                    print('\n''标准答案为：*************')
                    score6 = input('请输入第6题得分（10分）:')
                    while ((not (score6.isdigit())) or (int(score6) > 10) or (int(score6) < 0)):
                        score6 = input('请输入分值范围内的数字！第6题得分（10分）:')
                    print('第6题得分为：' + score6)
                    os.system("cls")

                    print(f'学生{k}：' + name + ID + '，共有图片' + str(pic_num) + '张')
                    ##===========================给出第七题得分==============================================
                    chongpi7 = ''
                    print("\n")
                    while (chongpi7 != '0'):
                        leixing = input('请输入第7题类型:')
                        print("\n")
                        while ((not (leixing.isdigit())) or (int(leixing) < 1) or (int(leixing) > 3)):
                            leixing = input('请输入1-3之间的类型数字代号:')
                        cell_ = [0]
                        if leixing == '1':
                            for i in range(1, 17, 1):
                                try:
                                    m = tb[1].rows[i].cells[2]
                                    cell_.append(m.text)
                                except:
                                    continue
                            score7 = str(check_1(cell_))
                        if leixing == '2':
                            for i in range(1, 17, 1):
                                try:
                                    m = tb[1].rows[i].cells[2]
                                    cell_.append(m.text)
                                except:
                                    continue
                            score7 = str(check_2(cell_))
                        if leixing == '3':
                            for i in range(1, 17, 1):
                                try:
                                    m = tb[1].rows[i].cells[2]
                                    cell_.append(m.text)
                                except:
                                    continue
                            score7 = str(check_3(cell_))

                        print('第 7 题学生答案:', cell_)
                        print("\n")

                        if int(score7) < 10:
                            score7 = '10'
                        print('第7题得分为：' + score7)
                        chongpi7 = input('输入0确认结果，其他重批:')
                        os.system("cls")

                    print(f'学生{k}：' + name + ID + '，共有图片' + str(pic_num) + '张')

                    ##===========================给出第八题得分==============================================
                    answer8 = ''
                    for l in range(1, 15, 1):
                        try:
                            cell = tb[2].rows[l].cells[2]
                            answer8 = answer8 + '\n' + cell.text
                        except:
                            continue
                    print('第 8 题学生答案:' + answer8)
                    print('\n''第 8 题标准答案：' + '*********************')
                    score8 = input('请输入第8题得分（20分）:')
                    while ((not (score8.isdigit())) or (int(score8) > 20) or (int(score8) < 0)):
                        score8 = input('请输入分值范围内的数字！第8题得分（20分）:')
                    print('第8题得分为：' + score8)
                    os.system("cls")

                    print(f'学生{k}：' + name + ID + '，共有图片' + str(pic_num) + '张')
                    ##===========================给出第九题得分==============================================
                    if len(answer9.strip()) > 1:
                        print("第 9 题回答是：" + answer9.strip())
                    else:
                        print("第 9 题回答可能是图片类型：" + answer9.strip())
                    print('\n''标准答案为：*************')
                    score9 = input('请输入第9题得分（10分）:')
                    while ((not (score9.isdigit())) or (int(score9) > 10) or (int(score9) < 0)):
                        score9 = input('请输入分值范围内的数字！第9题得分（10分）:')
                    print('第9题得分为：' + score9)
                    os.system("cls")

                    print(f'学生{k}：' + name + ID + '，共有图片' + str(pic_num) + '张')
                    ##===========================给出第十题得分==============================================
                    if len(answer10.strip()) > 1:
                        print("第 10 题回答是：" + answer10.strip())
                    else:
                        print("第 10 题回答可能是图片类型：" + answer10.strip())
                    print('\n''标准答案为：*************')
                    score10 = input('请输入第10题得分（10分）:')
                    while ((not (score10.isdigit())) or (int(score10) > 10) or (int(score10) < 0)):
                        score10 = input('请输入分值范围内的数字！第10题得分（10分）:')
                    print('第10题得分为：' + score10)
                    os.system("cls")
                    ##===========================计算分数==============================================
                    yuxiscore = str(int(score1) + int(score2) + int(score3) + int(score4) + int(score5))  # 计算预习部分成绩
                    guochengscore = str(12)
                    baogaoscore = str(int(score6) + int(score7) + int(score8) + int(score9) + int(score10))
                    zongscore = str(int(yuxiscore) + int(guochengscore) + int(baogaoscore))

                    print(f'学生{k}：' + name + ID + '，共有图片' + str(pic_num) + '张')
                    print(
                        '小成绩统计：' + score1 + ' ' + score2 + ' ' + score3 + ' ' + score4 + ' ' + score5 + ' ' + score6 + ' ' + score7 + ' ' + score8 + ' ' + score9 + ' ' + score10)
                    print('预习部分：' + yuxiscore)
                    print('实验过程表现：' + guochengscore)
                    print('实验报告部分：' + baogaoscore)
                    print('总成绩部分：' + zongscore)
                    time.sleep(2)
                    os.system("cls")

                    # ===========打印姓名学号到表中=============
                    cell = tb[0].rows[4].cells[1]
                    name = cell.text
                    ws.cell(row=krow, column=1, value=name)

                    cell = tb[0].rows[3].cells[4]
                    ID = cell.text
                    ws.cell(row=krow, column=2, value=ID)

                    # ==============打印成绩到表中===============
                    ws.cell(row=krow, column=3, value=score1)
                    ws.cell(row=krow, column=4, value=score2)
                    ws.cell(row=krow, column=5, value=score3)
                    ws.cell(row=krow, column=6, value=score4)
                    ws.cell(row=krow, column=7, value=score5)

                    ws.cell(row=krow, column=8, value=score6)
                    ws.cell(row=krow, column=9, value=score7)
                    ws.cell(row=krow, column=10, value=score8)
                    ws.cell(row=krow, column=11, value=score9)
                    ws.cell(row=krow, column=12, value=score10)

                    ws.cell(row=krow, column=13, value=yuxiscore)
                    ws.cell(row=krow, column=14, value=guochengscore)
                    ws.cell(row=krow, column=15, value=baogaoscore)
                    ws.cell(row=krow, column=16, value=zongscore)

                    # ===================================打印成绩到word文档里==================================

                    cell = tb[0].rows[7].cells[1]
                    run = cell.paragraphs[0]
                    run.clear()
                    run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 设置居中
                    run1 = run.add_run(yuxiscore)
                    run1.font.size = Pt(12)  # 设置字号小四
                    run1.font.color.rgb = RGBColor(255, 0, 0)  # 设置红色

                    cell = tb[0].rows[7].cells[2]
                    run = cell.paragraphs[0]
                    run.clear()
                    run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 设置居中
                    run1 = run.add_run(guochengscore)
                    run1.font.size = Pt(12)  # 设置字号小四
                    run1.font.color.rgb = RGBColor(255, 0, 0)  # 设置红色

                    cell = tb[0].rows[7].cells[3]
                    run = cell.paragraphs[0]
                    run.clear()
                    run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 设置居中
                    run1 = run.add_run(baogaoscore)
                    run1.font.size = Pt(12)  # 设置字号小四
                    run1.font.color.rgb = RGBColor(255, 0, 0)  # 设置红色

                    cell = tb[0].rows[7].cells[4]
                    run = cell.paragraphs[0]
                    run.clear()
                    run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 设置居中
                    run1 = run.add_run(zongscore)
                    run1.font.size = Pt(12)  # 设置字号小四
                    run1.font.color.rgb = RGBColor(255, 0, 0)  # 设置红色

                    # =========================================word内给出分数===========================================================
                    par1 = paragraphs[i1].add_run("                                      (  " + score1 + "  ) ")
                    par1.font.size = Pt(12)
                    par1.font.color.rgb = RGBColor(255, 0, 0)

                    par2 = paragraphs[i2].add_run(
                        "                                                         (  " + score2 + "  ) ")
                    par2.font.size = Pt(12)
                    par2.font.color.rgb = RGBColor(255, 0, 0)

                    par2 = paragraphs[i3].add_run("                                              (  " + score3 + "  ) ")
                    par2.font.size = Pt(12)
                    par2.font.color.rgb = RGBColor(255, 0, 0)

                    par2 = paragraphs[i4].add_run(
                        "                                                   (  " + score4 + "  ) ")
                    par2.font.size = Pt(12)
                    par2.font.color.rgb = RGBColor(255, 0, 0)

                    par2 = paragraphs[i5].add_run(
                        "                                                      (  " + score5 + "  ) ")
                    par2.font.size = Pt(12)
                    par2.font.color.rgb = RGBColor(255, 0, 0)

                    par2 = paragraphs[i6].add_run("                                (  " + score6 + "  ) ")
                    par2.font.size = Pt(12)
                    par2.font.color.rgb = RGBColor(255, 0, 0)

                    par2 = paragraphs[i7].add_run("                                         (  " + score7 + "  ) ")
                    par2.font.size = Pt(12)
                    par2.font.color.rgb = RGBColor(255, 0, 0)

                    par2 = paragraphs[i8].add_run("                                (  " + score8 + "  ) ")
                    par2.font.size = Pt(12)
                    par2.font.color.rgb = RGBColor(255, 0, 0)

                    par2 = paragraphs[i9].add_run("                   (  " + score9 + "  ) ")
                    par2.font.size = Pt(12)
                    par2.font.color.rgb = RGBColor(255, 0, 0)

                    par2 = paragraphs[i10].add_run("                          (  " + score10 + "  ) ")
                    par2.font.size = Pt(12)
                    par2.font.color.rgb = RGBColor(255, 0, 0)

                    banhao = tb[0].rows[3].cells[1].text

                    wb.save(newexcel)

                    document.save(os.path.join(pigaifilepath, word))
                else:
                    print('学生{k}：' + name + ID + '报告内容有异常（题目不全）,已经跳过')
                    print(str(n))
                    time.sleep(2)

                    yichang += 1
                    continue

            except:
                print('学生{k}：' + name + ID + '报告内容有异常（程序中断）,已经跳过')
                time.sleep(2)
                yichang += 1
                continue
        print('本班级已经批改完毕，有' + str(yichang) + '个内容异常报告，转格式失败的实验报告有 ' + str(ms) + ' 个。')

        excel1path = r'd:\减速器拆装与测量实验报告批改\减速器拆装与测量实验报告批改工作区\本班成绩表.xlsx'  # 批改程序自动产生的表
        excel2path = f'd:\减速器拆装与测量实验报告批改\减速器拆装与测量实验报告批改工作区\{class_name}\{class_name}.xlsx'  # 需要填写成绩的表（班级表）
        if (os.path.exists(
                f"D:\减速器拆装与测量实验报告批改\批完减速器拆装与测量实验报告作业\减速器拆装与测量实验报告各班成绩单汇总\{faculty}系")):
            excel3path = f"D:\减速器拆装与测量实验报告批改\批完减速器拆装与测量实验报告作业\减速器拆装与测量实验报告各班成绩单汇总\{faculty}系"
        else:
            excel3_third = os.mkdir(
                f"D:\减速器拆装与测量实验报告批改\批完减速器拆装与测量实验报告作业\减速器拆装与测量实验报告各班成绩单汇总\{faculty}系")
            excel3path = f"D:\减速器拆装与测量实验报告批改\批完减速器拆装与测量实验报告作业\减速器拆装与测量实验报告各班成绩单汇总\{faculty}系"

        shutil.copy(excel1path, doc_to_docx_path + f"\\{class_name}批改.xlsx")
        ##================================================

        ##======================激活表格================================
        wb1 = load_workbook(excel1path)  # 打开指定目录下的表格
        ws1 = wb1.active  # 激活表格

        wb2 = load_workbook(excel2path)  # 打开指定目录下的表格
        ws2 = wb2.active  # 激活表格
        ##=======识别原表有成绩的总行数，为下面程序提供遍历终止条件=====
        krow = 2  # 定义写入c成绩表格的初始行号
        kcolumn = 1  # 定义表格总行数
        kong = 0
        i = 0  # 标记循环次数，限制在101次
        while not (i == 100):
            i += 1
            if ws1.cell(row=krow, column=1).value is not None:  # 如果原表该行有成绩
                kcolumn += 1  # 总行数+1
                krow += 1  # 行数索引+1
            else:
                for j in range(1, 11):  # 判断当前空行的下十行内是否也为空
                    if (ws1.cell(row=krow + j, column=1).value is not None):
                        kong = 0
                        break
                    else:
                        kong = 1

                if kong == 0:  # 该行虽然没成绩，但下十行中有有成绩的行，认为是原表批改有漏掉的同学，总行数继续增加
                    kcolumn += 1  # 总行数+1
                    krow += 1  # 行数索引+1
                else:
                    oldrnum = kcolumn  # 原表总行数保存在oldrnum里
                    print('原表行数识别为:' + str(kcolumn))
                    break

        ##=======识别目标表格的总行数===================
        kong = 0
        krow = 2  # 定义写入c成绩表格的初始行号
        kcolumn = 1  # 定义表格总行数
        i = 0  # 标记循环次数，限制在101次
        while not (i == 100):
            i += 1
            if ws2.cell(row=krow, column=4).value is not None:  # 如果原表该行有成绩
                kcolumn += 1  # 总行数+1
                krow += 1  # 行数索引+1
            else:
                for j in range(1, 11):  # 判断当前空行的下十行内是否也为空
                    if (ws2.cell(row=krow + j, column=4).value is not None):
                        kong = 0
                        break
                    else:
                        kong = 1

                if kong == 0:  # 该行虽然没成绩，但下十行中有有成绩的行，认为是原表批改有漏掉的同学，总行数继续增加
                    kcolumn += 1  # 总行数+1
                    krow += 1  # 行数索引+1
                else:
                    newrnum = kcolumn  # 新表总行数保存在newrnum里
                    print('目标表行数识别为:' + str(kcolumn))
                    break

        ##=======将成绩填入新表===================
        k = 0  # 原表学生信息为空的数量
        for i in range(2, oldrnum + 1):
            name1 = ws1.cell(row=i, column=1).value
            if name1 is not None:
                yuxi = ws1.cell(row=i, column=13).value
                guocheng = ws1.cell(row=i, column=14).value
                baogao = ws1.cell(row=i, column=15).value
                zong = ws1.cell(row=i, column=16).value
                for j in range(2, newrnum + 1):
                    name2 = ws2.cell(row=j, column=4).value
                    if name2 == name1:
                        ws2.cell(row=j, column=5, value=yuxi)
                        ws2.cell(row=j, column=6, value=guocheng)
                        ws2.cell(row=j, column=7, value=baogao)
                        ws2.cell(row=j, column=8, value=zong)
                        break
                    else:
                        if j == newrnum:
                            print('目标表格中未找到学生：' + str(name1))
            else:
                k += 1

        if k == 0:
            wb2.save(excel2path)
            print('导入完成,已保存')
        else:
            wb2.save(excel2path)
            print(f'导入完成，已保存。但原表中有{k}个学生的信息为空')
        shutil.copy(excel2path, excel3path)

        shutil.move(excel3path + "\\" + f"{class_name}.xlsx", excel3path + "\\" + f"{class_name}成绩单.xlsx")
        print('\n'f'含有小分的本班成绩表保存在‘{class_name}批改.xlsx’表格内，已同批改后的作业放到一起。')
        print(f"{class_name}班成绩导入已经完成！")
        try:
            shutil.rmtree(wordfilepath)
        except:
            print("\n"f"无权限删除{class_name}的原文件夹，请手动删除！！！！！！！！！！！！！！！！！！！！！！！！！！！\n")
        judge = input("请输入是否继续批改下一个班级：\nY    or    N\n")
        if ((judge == "N") or (judge == "n")):
            break


if __name__ == "__main__":
    homework_correcting()